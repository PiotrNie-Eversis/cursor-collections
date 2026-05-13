"""Tests for chapter map and update (python-docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from chapters import build_sections, read_section_text, update_section_body


def test_build_sections_and_update(tmp_path: Path) -> None:
    doc = Document()
    h = doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Body one old.")
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("Body two old.")

    p = tmp_path / "t.docx"
    doc.save(str(p))

    d2 = Document(str(p))
    secs = build_sections(d2)
    assert [s.section_id for s in secs] == ["sec-0", "sec-1"]
    assert "Body one old." in read_section_text(d2, secs[0])
    assert "Body two old." in read_section_text(d2, secs[1])

    update_section_body(d2, secs[0], "Body one **new**.")
    d2.save(str(p))

    d3 = Document(str(p))
    secs3 = build_sections(d3)
    text = read_section_text(d3, secs3[0])
    assert "Body one **new**." in text
    assert "Body one old." not in text
    assert "Body two old." in read_section_text(d3, secs3[1])


def test_no_headings_single_section(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Only body.")
    p = tmp_path / "plain.docx"
    doc.save(str(p))
    d = Document(str(p))
    secs = build_sections(d)
    assert len(secs) == 1
    assert secs[0].section_id == "sec-0"
    assert "Only body." in read_section_text(d, secs[0])
