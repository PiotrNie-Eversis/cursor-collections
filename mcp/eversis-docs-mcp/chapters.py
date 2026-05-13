"""Outline and section boundaries for .docx files (python-docx).

``chapter_id`` format: ``sec-N`` where N is 0-based index of a section.
Section 0 is optional preamble before the first heading; thereafter each section
starts at a paragraph styled as a heading (style name contains ``Heading`` or ``heading``).
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """Insert a new paragraph after ``paragraph`` (same document body)."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _is_heading(p: Paragraph) -> bool:
    style = p.style and p.style.name
    if not style:
        return False
    s = style.lower()
    return "heading" in s or s.startswith("toc")


def _heading_level(p: Paragraph) -> int:
    """Best-effort heading level (1–9); 1 for unknown heading styles."""
    style = p.style and p.style.name
    if not style:
        return 1
    # Word often uses "Heading 1", "Heading 2"
    parts = style.replace("Heading", "").strip().split()
    for part in parts:
        if part.isdigit():
            return max(1, min(9, int(part)))
    return 1


@dataclass
class Section:
    section_id: str
    level: int
    title: str
    start: int  # paragraph index (inclusive)
    end: int  # paragraph index (exclusive)


def _paragraphs(doc: Document) -> list[Paragraph]:
    return list(doc.paragraphs)


def build_sections(doc: Document) -> list[Section]:
    paras = _paragraphs(doc)
    sections: list[Section] = []
    heading_indices: list[tuple[int, int, str]] = []  # idx, level, text

    for i, p in enumerate(paras):
        if _is_heading(p):
            t = (p.text or "").strip() or "(untitled)"
            heading_indices.append((i, _heading_level(p), t))

    if not heading_indices:
        # single body section
        return [Section(section_id="sec-0", level=1, title="Document", start=0, end=len(paras))]

    # Preamble before first heading
    first_heading_idx = heading_indices[0][0]
    if first_heading_idx > 0:
        sections.append(
            Section(
                section_id="sec-0",
                level=1,
                title="Preamble",
                start=0,
                end=first_heading_idx,
            )
        )

    for j, (idx, level, title) in enumerate(heading_indices):
        start = idx
        if j + 1 < len(heading_indices):
            end = heading_indices[j + 1][0]
        else:
            end = len(paras)
        sec_num = len(sections)
        sections.append(
            Section(
                section_id=f"sec-{sec_num}",
                level=level,
                title=title,
                start=start,
                end=end,
            )
        )

    return sections


def read_section_text(doc: Document, section: Section) -> str:
    paras = _paragraphs(doc)
    lines: list[str] = []
    for i in range(section.start, min(section.end, len(paras))):
        if i == section.start and _is_heading(paras[i]):
            continue
        t = (paras[i].text or "").strip()
        if t:
            lines.append(t)
    return "\n\n".join(lines)


def render_summary_md(sections: list[Section], doc_path: Path) -> str:
    lines = [
        f"<!-- Generated map for `{doc_path.name}` — do not edit structure by hand unless needed. -->",
        "",
        "| chapter_id | level | title |",
        "| --- | --- | --- |",
    ]
    for s in sections:
        title = s.title.replace("|", "\\|")
        lines.append(f"| `{s.section_id}` | {s.level} | {title} |")
    return "\n".join(lines)


def resolve_docx_path(docx_path: str) -> Path:
    p = Path(docx_path).expanduser()
    if not p.is_file():
        raise FileNotFoundError(f"Not a file: {p}")
    return p.resolve()


def load_sections(docx_path: str) -> tuple[Path, Document, list[Section]]:
    path = resolve_docx_path(docx_path)
    doc = Document(str(path))
    return path, doc, build_sections(doc)


def update_section_body(
    doc: Document,
    section: Section,
    new_content: str,
    *,
    placeholder_graphics: str | None = None,
) -> None:
    """Replace non-heading paragraphs in section (excluding title line) with new body text."""
    paras = _paragraphs(doc)
    if section.start >= len(paras):
        return
    start_body = (
        section.start + 1 if _is_heading(paras[section.start]) else section.start
    )
    end = min(section.end, len(paras))

    to_remove: list = []
    for i in range(start_body, end):
        to_remove.append(paras[i]._element)
    for el in reversed(to_remove):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    paras = _paragraphs(doc)
    if section.start >= len(paras):
        return
    anchor = paras[section.start]

    body = new_content.rstrip()
    if placeholder_graphics:
        body = f"{body}\n\n{placeholder_graphics.strip()}"

    first_block = True
    for line in body.split("\n\n"):
        line = line.strip()
        if not line:
            continue
        if first_block:
            anchor = _insert_paragraph_after(anchor, line)
            first_block = False
        else:
            anchor = _insert_paragraph_after(anchor, line)
